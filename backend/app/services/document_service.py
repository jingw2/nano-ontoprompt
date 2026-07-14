import os
from dataclasses import dataclass

CONVERSION_ERROR_MARKERS = (
    "[File conversion failed:",
    "[Text read failed:",
    "[CSV read failed:",
)


@dataclass
class ConversionResult:
    content: str | None = None
    error: str | None = None

    @property
    def ok(self) -> bool:
        return self.error is None and bool(self.content and self.content.strip())


def is_usable_converted_text(text: str | None) -> bool:
    if not text or not text.strip():
        return False
    stripped = text.strip()
    return not any(stripped.startswith(marker) for marker in CONVERSION_ERROR_MARKERS)


def combine_converted_files(files) -> tuple[str | None, str | None]:
    """Return (combined_text, error_message). error_message is set when no usable content."""
    usable = [f for f in files if is_usable_converted_text(getattr(f, "converted_md", None))]
    if not usable:
        bad_names = [getattr(f, "filename", "?") for f in files]
        return None, f"以下文件无法用于提取（转换失败或无文本内容）：{', '.join(bad_names)}"

    combined = "\n\n---\n\n".join(
        f"【来源文件】{getattr(f, 'filename', 'unknown')}\n\n{f.converted_md.strip()}"
        for f in usable
    )
    if not combined.strip():
        return None, "上传的文件中没有可用于提取的文本内容"
    return combined, None


def convert_document(file_path: str, mime_type: str | None = None) -> ConversionResult:
    ext = os.path.splitext(file_path)[1].lower()

    if ext in (".md", ".txt") or (mime_type and ("text/plain" in mime_type or "text/markdown" in mime_type)):
        return _read_plain_text(file_path)

    if ext == ".csv" or (mime_type and "csv" in mime_type):
        return _read_csv_as_markdown(file_path)

    docx_result: ConversionResult | None = None
    if ext == ".docx":
        docx_result = _convert_docx(file_path)
        if docx_result.ok:
            return docx_result

    try:
        from markitdown import MarkItDown

        md = MarkItDown()
        result = md.convert(file_path)
        content = (result.text_content or "").strip()
        if content:
            return ConversionResult(content=content)
        return ConversionResult(error="文件转换后没有可用文本内容")
    except BaseException as e:
        if docx_result is not None:
            return docx_result
        if mime_type and ("text" in mime_type or "csv" in mime_type):
            text_result = _read_plain_text(file_path)
            if text_result.ok:
                return text_result
        return ConversionResult(error=f"文件转换失败：{e}")


def _read_plain_text(file_path: str) -> ConversionResult:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read().strip()
        if not content:
            return ConversionResult(error="文件为空")
        return ConversionResult(content=content)
    except Exception as e:
        return ConversionResult(error=f"文本读取失败：{e}")


def _read_csv_as_markdown(file_path: str) -> ConversionResult:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            lines = f.read().splitlines()
        if not lines:
            return ConversionResult(error="CSV 文件为空")
        header = lines[0]
        separator = "|".join(["---"] * (header.count(",") + 1))
        md_lines = [
            "| " + header.replace(",", " | ") + " |",
            "| " + separator + " |",
        ]
        for row in lines[1:]:
            md_lines.append("| " + row.replace(",", " | ") + " |")
        return ConversionResult(content="\n".join(md_lines))
    except Exception as e:
        return ConversionResult(error=f"CSV 读取失败：{e}")


def _convert_docx(file_path: str) -> ConversionResult:
    try:
        from docx import Document

        doc = Document(file_path)
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append("| " + " | ".join(cells) + " |")
            if rows:
                if len(rows) > 1:
                    col_count = rows[0].count("|") - 1
                    sep = "| " + " | ".join(["---"] * max(col_count, 1)) + " |"
                    parts.append("\n".join([rows[0], sep, *rows[1:]]))
                else:
                    parts.append("\n".join(rows))

        content = "\n\n".join(parts).strip()
        if not content:
            return ConversionResult(error="DOCX 文件中没有可提取的文本")
        return ConversionResult(content=content)
    except Exception as e:
        return ConversionResult(error=f"DOCX 解析失败：{e}")
